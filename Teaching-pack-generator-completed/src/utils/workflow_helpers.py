"""
Workflow Helper Functions
Separate business logic from main workflow orchestration
"""
import os
import json
import random
import csv
import re
from datetime import datetime
from typing import List, Tuple, Optional, Dict
import pandas as pd
from loguru import logger
from pathlib import Path

from models.teaching_pack_models import (
    StudentDiagnosticResult, Diagnostic, SkillSet, LessonSummary
)
from utils.basetools.pdf_parser import extract_text_from_pdf, read_text_file

try:
    from docx import Document
    from docx.shared import Inches
except ImportError:
    Document = None


def load_lesson_content(file_path: str) -> str:
    """Load lesson content from PDF or TXT file"""
    if file_path.endswith('.pdf'):
        return extract_text_from_pdf(file_path)
    elif os.path.isfile(file_path):
        return read_text_file(file_path)
    else:
        return file_path


def parse_student_list_with_scores(file_path: str) -> List[Dict]:
    """
    Parse student list WITH SCORES from various file formats
    Expected columns: Name/Student, Toan, Van, Anh, etc.
    
    Args:
        file_path: Path to student list file
    
    Returns:
        List of student dicts: [{"student_id": "...", "full_name": "...", "subject_scores": {"Toan": 8.5, ...}}]
    
    Raises:
        ValueError: If file format is not supported or parsing fails
    """
    file_path_obj = Path(file_path)
    file_ext = file_path_obj.suffix.lower()
    
    try:
        # Excel or CSV files
        if file_ext in ['.xlsx', '.xls', '.xlsm', '.csv']:
            if file_ext == '.csv':
                df = pd.read_csv(file_path)
            else:
                df = pd.read_excel(file_path)
            
            # Find student name column
            name_col = None
            for col in df.columns:
                col_lower = str(col).lower()
                if any(keyword in col_lower for keyword in ['name', 'student', 'tn', 'hc sinh', 'h tn']):
                    name_col = col
                    break
            
            if not name_col:
                name_col = df.columns[0]  # Fallback to first column
            
            # Find student ID column (if exists)
            id_col = None
            for col in df.columns:
                col_lower = str(col).lower()
                if any(keyword in col_lower for keyword in ['id', 'mssv', 'ma', 'student_id']):
                    id_col = col
                    break
            
            # Find grade level column (qualitative assessment)
            grade_col = None
            for col in df.columns:
                col_lower = str(col).lower()
                if any(keyword in col_lower for keyword in ['grade', 'level', 'trnh ', 'nh gi', 'xp loi']):
                    grade_col = col
                    break
            
            # Find notes column
            notes_col = None
            for col in df.columns:
                col_lower = str(col).lower()
                if any(keyword in col_lower for keyword in ['notes', 'ghi ch', 'note', 'comment', 'nhn xt']):
                    notes_col = col
                    break
            
            # Find score columns (anything that's numeric and not ID)
            score_columns = []
            for col in df.columns:
                if col != name_col and col != id_col and pd.api.types.is_numeric_dtype(df[col]):
                    score_columns.append(col)
            
            # Parse students
            students_dict = {}
            for idx, row in df.iterrows():
                student_name = str(row[name_col]).strip()
                if not student_name or student_name.lower() == 'nan':
                    continue
                
                student_id = str(row[id_col]).strip() if id_col else student_name
                
                # Skip if student_id already exists
                if student_id in students_dict:
                    logger.warning(f"Duplicate student_id {student_id}, skipping")
                    continue
                
                # Parse scores
                subject_scores = {}
                for col in score_columns:
                    try:
                        score = float(row[col])
                        if not pd.isna(score):
                            subject_scores[col] = score
                    except:
                        pass
                
                # Get qualitative data
                grade_level = str(row[grade_col]).strip() if grade_col and grade_col in row and not pd.isna(row[grade_col]) else None
                notes = str(row[notes_col]).strip() if notes_col and notes_col in row and not pd.isna(row[notes_col]) else None
                
                students_dict[student_id] = {
                    "student_id": student_id,
                    "full_name": student_name,
                    "subject_scores": subject_scores,
                    "grade_level": grade_level,
                    "notes": notes
                }
            
            students = list(students_dict.values())
            logger.info(f"Loaded {len(students)} students with scores from file")
            return students
        
        # JSON files
        elif file_ext == '.json':
            with open(file_path, 'r', encoding='utf-8') as f:
                data = json.load(f)
            
            students = []
            if isinstance(data, list):
                for item in data:
                    if isinstance(item, dict):
                        students.append({
                            "student_id": item.get("student_id", item.get("id", item.get("name", ""))),
                            "full_name": item.get("full_name", item.get("name", "")),
                            "subject_scores": item.get("subject_scores", item.get("scores", {})),
                            "grade_level": item.get("grade_level"),
                            "notes": item.get("notes"),
                            "email": item.get("email")  # Only use email if explicitly provided
                        })
            
            logger.info(f"Loaded {len(students)} students from JSON file")
            return students
        
        else:
            raise ValueError(f"Unsupported file format: {file_ext}. For scores, use Excel or CSV.")
    
    except Exception as e:
        logger.error(f"Failed to parse student list with scores: {str(e)}")
        raise ValueError(f"Failed to parse student list with scores: {str(e)}")


def parse_student_list_file(file_path: str) -> List[str]:
    """
    Parse student list from various file formats (Excel, CSV, TXT, JSON)
    Auto-detects format and extracts student names/IDs
    
    Args:
        file_path: Path to student list file
    
    Returns:
        List of student names/IDs
    
    Raises:
        ValueError: If file format is not supported or parsing fails
    """
    file_path_obj = Path(file_path)
    file_ext = file_path_obj.suffix.lower()
    
    try:
        # Excel files (.xlsx, .xls)
        if file_ext in ['.xlsx', '.xls', '.xlsm']:
            df = pd.read_excel(file_path)
            # Try to find student name column
            for col in df.columns:
                col_lower = str(col).lower()
                if any(keyword in col_lower for keyword in ['name', 'student', 'tn', 'hc sinh', 'h tn']):
                    student_list = df[col].astype(str).dropna().tolist()
                    break
            else:
                # If no matching column, use first column
                student_list = df.iloc[:, 0].astype(str).dropna().tolist()
            
            # Clean up student names
            student_list = [s.strip() for s in student_list if s and s.strip() and s.strip().lower() != 'nan']
            logger.info(f"Loaded {len(student_list)} students from Excel file")
            return student_list
        
        # CSV files
        elif file_ext == '.csv':
            with open(file_path, 'r', encoding='utf-8-sig') as f:
                # Try to detect delimiter
                sample = f.read(1024)
                f.seek(0)
                sniffer = csv.Sniffer()
                try:
                    dialect = sniffer.sniff(sample)
                    delimiter = dialect.delimiter
                except:
                    delimiter = ','
                
                reader = csv.reader(f, delimiter=delimiter)
                rows = list(reader)
                
                if not rows:
                    raise ValueError("CSV file is empty")
                
                # Check if first row is header
                header = rows[0]
                data_rows = rows[1:] if len(rows) > 1 else rows
                
                # Try to find student name column
                student_col_idx = 0
                for idx, col_name in enumerate(header):
                    col_lower = str(col_name).lower()
                    if any(keyword in col_lower for keyword in ['name', 'student', 'tn', 'hc sinh', 'h tn']):
                        student_col_idx = idx
                        break
                
                # Extract student names
                student_list = [row[student_col_idx].strip() for row in data_rows if row and len(row) > student_col_idx and row[student_col_idx].strip()]
                logger.info(f"Loaded {len(student_list)} students from CSV file")
                return student_list
        
        # Text files (.txt)
        elif file_ext == '.txt':
            with open(file_path, 'r', encoding='utf-8') as f:
                content = f.read()
            
            # Try different parsing strategies
            # 1. Line-by-line (most common)
            lines = [line.strip() for line in content.split('\n') if line.strip()]
            
            # 2. If lines contain numbers or bullets, extract names
            student_list = []
            for line in lines:
                # Remove numbering like "1. Name", "1) Name", "- Name", " Name"
                cleaned = re.sub(r'^[\d]+[\.\)\-\s]+', '', line)
                cleaned = re.sub(r'^[\-\\*\+]\s+', '', cleaned)
                cleaned = cleaned.strip()
                
                if cleaned and not cleaned.lower().startswith(('danh sch', 'list', 'students', 'class')):
                    student_list.append(cleaned)
            
            logger.info(f"Loaded {len(student_list)} students from TXT file")
            return student_list
        
        # JSON files
        elif file_ext == '.json':
            with open(file_path, 'r', encoding='utf-8') as f:
                data = json.load(f)
            
            student_list = []
            # Handle different JSON structures
            if isinstance(data, list):
                for item in data:
                    if isinstance(item, str):
                        student_list.append(item)
                    elif isinstance(item, dict):
                        # Try to find name field
                        for key in ['name', 'student_name', 'full_name', 'student', 'tn', 'h tn']:
                            if key in item:
                                student_list.append(str(item[key]))
                                break
            elif isinstance(data, dict):
                # Check if there's a students array
                for key in ['students', 'student_list', 'names', 'danh_sach']:
                    if key in data and isinstance(data[key], list):
                        for item in data[key]:
                            if isinstance(item, str):
                                student_list.append(item)
                            elif isinstance(item, dict):
                                for name_key in ['name', 'student_name', 'full_name', 'student', 'tn', 'h tn']:
                                    if name_key in item:
                                        student_list.append(str(item[name_key]))
                                        break
                        break
            
            logger.info(f"Loaded {len(student_list)} students from JSON file")
            return student_list
        
        else:
            raise ValueError(f"Unsupported file format: {file_ext}. Supported formats: .xlsx, .xls, .csv, .txt, .json")
    
    except Exception as e:
        logger.error(f"Failed to parse student list file: {str(e)}")
        raise ValueError(f"Failed to parse student list file: {str(e)}")


def load_student_list(file_path: Optional[str] = None, default_count: int = 30) -> List[str]:
    """
    Load student list from file or generate default list
    
    Args:
        file_path: Path to student list file (optional)
        default_count: Number of default students if no file provided
    
    Returns:
        List of student IDs/names
    """
    if file_path:
        try:
            student_list = parse_student_list_file(file_path)
            if student_list:
                return student_list
            else:
                logger.warning("Parsed student list is empty, using default")
        except Exception as e:
            logger.warning(f"Could not read student file: {e}, using default")
    
    # Return default list
    return [f"Student_{i+1}" for i in range(default_count)]


def generate_mock_diagnostic_results(
    student_list: List[str],
    diagnostic: Diagnostic,
    skill_set: SkillSet
) -> List[StudentDiagnosticResult]:
    """
    Generate mock diagnostic results for testing
    
    Args:
        student_list: List of student IDs
        diagnostic: Diagnostic questionnaire
        skill_set: Skill set for the lesson
    
    Returns:
        List of student diagnostic results
    """
    results = []
    num_questions = len(diagnostic.questions)
    if num_questions == 0:
        logger.warning("Diagnostic has 0 questions; generating empty answers with 0.0 scores.")
    
    for student_id in student_list:
        answers = {}
        correct_count = 0
        skill_mastery = {skill.skill_id: 0.0 for skill in skill_set.skills}
        
        for q in diagnostic.questions:
            # Simulate varying performance (70% success rate)
            is_correct = random.random() > 0.3
            
            if is_correct:
                answers[q.question_id] = q.correct_answer
                correct_count += 1
                if q.skill_id in skill_mastery:
                    skill_mastery[q.skill_id] = min(1.0, skill_mastery[q.skill_id] + 0.5)
            else:
                # Pick random wrong answer
                wrong_options = [opt for opt in q.options if opt != q.correct_answer]
                if wrong_options:
                     answers[q.question_id] = random.choice(wrong_options)
                else:
                     answers[q.question_id] = "N/A"
        
        student_result = StudentDiagnosticResult(
            student_id=student_id,
            student_name=student_id,
            answers=answers,
            score=(correct_count / num_questions) if num_questions else 0.0,
            skill_mastery=skill_mastery,
            misconceptions=[]
        )
        results.append(student_result)
    
    logger.info(f"Generated {len(results)} mock diagnostic results")
    return results


def export_diagnostic_results(
    lesson_summary: LessonSummary,
    skill_set: SkillSet,
    diagnostic: Diagnostic,
    output_dir: str = "outputs",
) -> str:
    """
    Export diagnostic results to JSON file
    
    Returns:
        Path to exported file
    """
    timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
    output_file = f"diagnostic_{timestamp}.json"
    
    export_data = {
        "lesson_summary": lesson_summary.model_dump(),
        "skill_set": skill_set.model_dump(),
        "diagnostic": diagnostic.model_dump()
    }
    
    os.makedirs(output_dir, exist_ok=True)
    
    output_path = Path(output_dir) / output_file
    with open(output_path, 'w', encoding='utf-8') as f:
        json.dump(export_data, f, ensure_ascii=False, indent=2)
    
    logger.info(f"Results exported to {output_path}")
    return output_file


def format_diagnostic_questions(diagnostic: Diagnostic) -> str:
    """
    Format diagnostic questions for display
    
    Returns:
        Formatted markdown string
    """
    questions_text = "\n\n".join([
        f"**Question {i+1}:** {q.question_text}\n"
        f"- A) {q.options[0]}\n"
        f"- B) {q.options[1]}\n"
        f"- C) {q.options[2]}\n"
        f"- D) {q.options[3]}\n"
        f"*Correct: {q.correct_answer}*"
        for i, q in enumerate(diagnostic.questions)
    ])
    
    return f"##  Diagnostic Questions:\n\n{questions_text}"


def export_final_results(
    lesson_summary: LessonSummary,
    skill_set: SkillSet,
    diagnostic: Diagnostic,
    labeled_groups: List,
    teaching_packs: List,
    student_count: int,
    output_dir: str = "outputs",
) -> str:
    """
    Export complete teaching pack results to JSON
    
    Returns:
        Path to exported file
    """
    timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
    output_file = f"teaching_packs_{timestamp}.json"
    
    final_data = {
        "lesson_summary": lesson_summary.model_dump(),
        "skill_set": skill_set.model_dump(),
        "diagnostic": diagnostic.model_dump(),
        "groups": [g.model_dump() for g in labeled_groups],
        "teaching_packs": teaching_packs,
        "student_count": student_count
    }
    
    os.makedirs(output_dir, exist_ok=True)
    
    output_path = Path(output_dir) / output_file
    with open(output_path, 'w', encoding='utf-8') as f:
        json.dump(final_data, f, ensure_ascii=False, indent=2)
    
    logger.info(f"Final results exported to {output_path}")
    return output_file


def export_quiz_to_word(quiz_data: Dict, lesson_title: str = "Quiz", output_dir: str = "outputs") -> str:
    """
    Export quiz data to Word document
    
    Args:
        quiz_data: Dictionary containing quiz information with questions
        lesson_title: Title for the quiz document
        
    Returns:
        Path to exported Word file
    """
    if Document is None:
        raise ImportError("python-docx is required for Word export")
    
    timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
    output_file = f"quiz_{timestamp}.docx"
    output_path = Path(output_dir) / output_file
    
    doc = Document()
    doc.add_heading(f'{lesson_title} - Quiz', 0)
    
    # Add quiz metadata if available
    if 'total_questions' in quiz_data:
        doc.add_paragraph(f'Total Questions: {quiz_data["total_questions"]}')
    if 'estimated_time' in quiz_data:
        doc.add_paragraph(f'Estimated Time: {quiz_data["estimated_time"]} minutes')
    
    doc.add_paragraph('')  # Empty line
    
    # Get questions
    questions = quiz_data.get('questions', [])
    if not questions and isinstance(quiz_data, list):
        questions = quiz_data
    
    for i, question in enumerate(questions, 1):
        # Question text
        question_text = question.get('question_text', f'Question {i}')
        doc.add_heading(f'Question {i}: {question_text}', level=2)
        
        # Question metadata
        metadata = []
        if question.get('question_id'):
            metadata.append(f'ID: {question["question_id"]}')
        if question.get('skill_id'):
            metadata.append(f'Skill: {question["skill_id"]}')
        if question.get('difficulty'):
            metadata.append(f'Difficulty: {question["difficulty"]}')
        
        if metadata:
            doc.add_paragraph(', '.join(metadata), style='Intense Quote')
        
        # Options (for multiple choice)
        options = question.get('options', [])
        if options:
            for j, option in enumerate(options):
                correct_marker = ' ' if option == question.get('correct_answer') else ''
                doc.add_paragraph(f'{chr(65 + j)}. {option}{correct_marker}')
        else:
            doc.add_paragraph('Open-ended question')
        
        # Hint
        if question.get('hint'):
            doc.add_paragraph(f'Hint: {question["hint"]}', style='Intense Quote')
        
        # Explanation
        if question.get('explanation'):
            doc.add_paragraph(f'Explanation: {question["explanation"]}', style='Intense Quote')
        
        doc.add_paragraph('')  # Empty line between questions
    
    os.makedirs(output_dir, exist_ok=True)
    doc.save(output_path)
    
    logger.info(f"Quiz exported to Word document: {output_path}")
    return output_file
